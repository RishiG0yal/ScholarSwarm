"""
DOCX Parser — Extracts text from Word documents with paragraph-level coordinate mapping.
Uses python-docx to iterate paragraphs. Page numbers are estimated since DOCX
doesn't store explicit page breaks in a reliable way.
"""

import os
import logging
from docx import Document
from backend.models.schemas import TextChunk, ParsedDocument

logger = logging.getLogger("rag_pipeline.parsers.docx")

# Rough estimate: average paragraphs per page for page estimation
_ESTIMATED_PARAS_PER_PAGE = 6


def parse_docx(file_path: str, file_id: str) -> ParsedDocument:
    """
    Parse a DOCX file and return structured text chunks with coordinates.

    Args:
        file_path: Absolute path to the DOCX file.
        file_id: Unique identifier for this file upload.

    Returns:
        ParsedDocument with text chunks mapped to estimated [Page X, Paragraph Y].
    """
    filename = os.path.basename(file_path)
    logger.info(f"Parsing DOCX: {filename}")

    doc = Document(file_path)
    chunks: list[TextChunk] = []
    chunk_counter = 0
    page_num = 1
    para_on_page = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check for explicit page breaks in paragraph runs
        has_page_break = False
        for run in para.runs:
            # Check the XML for page breaks
            if run._element.xml and "w:br" in run._element.xml and 'w:type="page"' in run._element.xml:
                has_page_break = True
                break

        if has_page_break:
            page_num += 1
            para_on_page = 0

        para_on_page += 1
        chunk_counter += 1

        # Fallback page estimation if no explicit breaks detected
        estimated_page = page_num
        if not has_page_break and para_on_page > _ESTIMATED_PARAS_PER_PAGE:
            page_num += 1
            para_on_page = 1
            estimated_page = page_num

        chunks.append(
            TextChunk(
                text=text,
                page_number=estimated_page,
                paragraph_number=para_on_page,
                source_file=filename,
                chunk_id=f"{file_id}_p{estimated_page}_para{para_on_page}",
            )
        )

    total_pages = page_num
    logger.info(f"  Extracted {len(chunks)} chunks (~{total_pages} estimated pages)")

    return ParsedDocument(
        file_id=file_id,
        filename=filename,
        file_type="docx",
        total_pages=total_pages,
        total_chunks=len(chunks),
        chunks=chunks,
    )
